# -*- coding: utf-8 -*-
"""
Comprehensive Platform for Explaining Original Econometric Papers and Tests
Design: Dr. Merwan Roudane
"""

import streamlit as st
from google import genai
from google.genai import types
import tempfile
import os
from datetime import datetime
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io

# Page configuration
st.set_page_config(
    page_title="Comprehensive Platform - Econometrics",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS for warm colors and professional design
st.markdown("""
<style>
    /* Main background - warm light colors */
    .stApp {
        background: linear-gradient(135deg, #FFFBF5 0%, #FFF8F0 50%, #FFF5EB 100%);
    }
    
    /* ========== Move sidebar to left-to-right layout ========== */
    /* Full LTR layout */
    [data-testid="stAppViewContainer"] {
        direction: ltr;
    }
    
    /* Inner content direction */
    [data-testid="stAppViewContainer"] > * {
        direction: ltr;
    }
    
    /* Sidebar direction */
    [data-testid="stSidebar"] {
        direction: ltr;
    }
    
    /* Sidebar background */
    [data-testid="stSidebar"] > div:first-child {
        background: linear-gradient(180deg, #FF9F6B 0%, #FF8C5A 100%) !important;
    }
    
    section[data-testid="stSidebar"] {
        background: linear-gradient(180deg, #FF9F6B 0%, #FF8C5A 100%) !important;
    }
    
    /* ========== Fix all white elements in sidebar ========== */
    
    /* Remove white background from all elements */
    [data-testid="stSidebar"] div,
    [data-testid="stSidebar"] section,
    [data-testid="stSidebar"] [data-testid="stVerticalBlock"],
    [data-testid="stSidebar"] [data-testid="stVerticalBlockBorderWrapper"] {
        background-color: transparent !important;
        background: transparent !important;
    }
    
    /* White text */
    [data-testid="stSidebar"] * {
        color: white !important;
    }
    
    [data-testid="stSidebar"] .stTextInput label,
    [data-testid="stSidebar"] .stFileUploader label,
    [data-testid="stSidebar"] label {
        color: white !important;
        font-weight: bold;
    }
    
    /* ========== API input field ========== */
    [data-testid="stSidebar"] input[type="password"],
    [data-testid="stSidebar"] input[type="text"],
    [data-testid="stSidebar"] .stTextInput input {
        background-color: rgba(255, 255, 255, 0.2) !important;
        color: white !important;
        border: 2px solid rgba(255, 255, 255, 0.4) !important;
        border-radius: 8px !important;
    }
    
    [data-testid="stSidebar"] input::placeholder {
        color: rgba(255, 255, 255, 0.7) !important;
    }
    
    /* ========== Dropdown select ========== */
    [data-testid="stSidebar"] [data-baseweb="select"] {
        background-color: rgba(255, 255, 255, 0.2) !important;
        border: 2px solid rgba(255, 255, 255, 0.4) !important;
        border-radius: 8px !important;
    }
    
    [data-testid="stSidebar"] [data-baseweb="select"] > div {
        background-color: transparent !important;
    }
    
    [data-testid="stSidebar"] .stSelectbox span,
    [data-testid="stSidebar"] [data-baseweb="select"] span {
        color: white !important;
    }
    
    /* ========== File upload area - fix white color ========== */
    [data-testid="stSidebar"] [data-testid="stFileUploader"],
    [data-testid="stSidebar"] [data-testid="stFileUploader"] > div,
    [data-testid="stSidebar"] [data-testid="stFileUploader"] section,
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"],
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"],
    [data-testid="stSidebar"] .uploadedFile {
        background-color: rgba(255, 255, 255, 0.15) !important;
        background: rgba(255, 255, 255, 0.15) !important;
        border-color: rgba(255, 255, 255, 0.4) !important;
    }
    
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] {
        border: 2px dashed rgba(255, 255, 255, 0.5) !important;
        border-radius: 12px !important;
        background: rgba(255, 255, 255, 0.1) !important;
    }
    
    /* Fix white box inside file uploader */
    [data-testid="stSidebar"] [data-testid="stFileUploader"] > div > div,
    [data-testid="stSidebar"] .stFileUploader > div > div,
    [data-testid="stSidebar"] [data-testid="stFileUploader"] [data-testid="stMarkdownContainer"],
    [data-testid="stSidebar"] [data-testid="stFileUploader"] small,
    [data-testid="stSidebar"] [data-testid="stFileUploader"] p {
        background: transparent !important;
        background-color: transparent !important;
    }
    
    /* Inner white area */
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"] > div,
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzoneInstructions"] > div > div {
        background: transparent !important;
        background-color: transparent !important;
    }
    
    /* Any div with white background */
    [data-testid="stSidebar"] div[style*="background-color: white"],
    [data-testid="stSidebar"] div[style*="background: white"],
    [data-testid="stSidebar"] div[style*="background-color: rgb(255, 255, 255)"],
    [data-testid="stSidebar"] div[style*="background: rgb(255, 255, 255)"] {
        background: transparent !important;
        background-color: transparent !important;
    }
    
    /* Browse files button */
    [data-testid="stSidebar"] [data-testid="stFileUploaderDropzone"] button,
    [data-testid="stSidebar"] [data-testid="baseButton-secondary"] {
        background-color: rgba(255, 255, 255, 0.25) !important;
        border: 1px solid rgba(255, 255, 255, 0.5) !important;
        color: white !important;
    }
    
    /* Uploaded files */
    [data-testid="stSidebar"] [data-testid="stFileUploaderFile"],
    [data-testid="stSidebar"] .uploadedFileData {
        background: rgba(255, 255, 255, 0.2) !important;
        border-radius: 8px !important;
    }
    
    /* ========== Expander (uploaded files) ========== */
    [data-testid="stSidebar"] [data-testid="stExpander"],
    [data-testid="stSidebar"] details,
    [data-testid="stSidebar"] summary {
        background: rgba(255, 255, 255, 0.15) !important;
        border: 1px solid rgba(255, 255, 255, 0.3) !important;
        border-radius: 10px !important;
    }
    
    [data-testid="stSidebar"] details > div {
        background: transparent !important;
    }
    
    /* ========== Dropdown menu options ========== */
    [data-baseweb="popover"],
    [data-baseweb="menu"] {
        direction: ltr;
    }
    
    [data-baseweb="popover"] li,
    [data-baseweb="menu"] li {
        text-align: left !important;
        color: #333 !important;
    }
    
    /* ========== Headings ========== */
    h1 {
        color: #E07B39 !important;
        text-align: center;
        font-weight: bold;
        text-shadow: 1px 1px 2px rgba(0,0,0,0.1);
    }
    
    h2, h3 {
        color: #D4703A !important;
        text-align: left;
        direction: ltr;
    }
    
    /* ========== LTR for main content text ========== */
    .main p, 
    .main div,
    .main span,
    .main li,
    .main h1, .main h2, .main h3, .main h4, .main h5, .main h6,
    [data-testid="stMarkdownContainer"] p,
    [data-testid="stMarkdownContainer"] div,
    [data-testid="stMarkdownContainer"] span,
    [data-testid="stMarkdownContainer"] li,
    [data-testid="stMarkdownContainer"] h1,
    [data-testid="stMarkdownContainer"] h2,
    [data-testid="stMarkdownContainer"] h3,
    [data-testid="stMarkdownContainer"] h4,
    [data-testid="stMarkdownContainer"] ul,
    [data-testid="stMarkdownContainer"] ol {
        direction: ltr !important;
        text-align: left !important;
        unicode-bidi: plaintext;
    }
    
    /* Lists */
    .main ul, .main ol,
    [data-testid="stMarkdownContainer"] ul,
    [data-testid="stMarkdownContainer"] ol {
        direction: ltr !important;
        text-align: left !important;
        padding-left: 20px !important;
        padding-right: 0 !important;
    }
    
    /* List items */
    .main li,
    [data-testid="stMarkdownContainer"] li {
        direction: ltr !important;
        text-align: left !important;
    }
    
    /* Tables */
    .main table,
    [data-testid="stMarkdownContainer"] table {
        direction: ltr !important;
    }
    
    .main th, .main td,
    [data-testid="stMarkdownContainer"] th,
    [data-testid="stMarkdownContainer"] td {
        text-align: left !important;
    }
    
    /* Code blocks stay LTR */
    .main code, .main pre,
    [data-testid="stMarkdownContainer"] code,
    [data-testid="stMarkdownContainer"] pre {
        direction: ltr !important;
        text-align: left !important;
        unicode-bidi: embed;
    }
    
    /* ========== Buttons ========== */
    .stButton > button {
        background: linear-gradient(90deg, #FF8C5A 0%, #FFAB76 100%);
        color: white !important;
        border: none;
        border-radius: 12px;
        padding: 12px 28px;
        font-weight: bold;
        font-size: 16px;
        transition: all 0.3s ease;
        box-shadow: 0 4px 15px rgba(255, 140, 90, 0.3);
    }
    
    .stButton > button:hover {
        background: linear-gradient(90deg, #E07B39 0%, #FF8C5A 100%);
        transform: translateY(-3px);
        box-shadow: 0 6px 20px rgba(255, 140, 90, 0.4);
    }
    
    /* Sidebar buttons */
    [data-testid="stSidebar"] .stButton > button {
        background: rgba(255, 255, 255, 0.25) !important;
        border: 1px solid rgba(255, 255, 255, 0.5) !important;
    }
    
    [data-testid="stSidebar"] .stButton > button:hover {
        background: rgba(255, 255, 255, 0.35) !important;
    }
    
    /* Cards */
    .info-card {
        background: white;
        border-radius: 16px;
        padding: 24px;
        box-shadow: 0 4px 20px rgba(255, 140, 90, 0.12);
        border: 1px solid #FFE4D6;
        margin: 12px 0;
        direction: ltr;
        text-align: left;
    }
    
    /* Chat messages */
    .chat-user {
        background: linear-gradient(135deg, #FFE8D6 0%, #FFD9C0 100%);
        border-radius: 18px 18px 5px 18px;
        padding: 16px 20px;
        margin: 12px 0;
        border-left: 4px solid #FF8C5A;
        border-right: none;
        animation: slideInLeft 0.4s ease;
        direction: ltr;
        text-align: left;
    }
    
    .chat-assistant {
        background: linear-gradient(135deg, #FFFFFF 0%, #FFF8F0 100%);
        border-radius: 18px 18px 18px 5px;
        padding: 16px 20px;
        margin: 12px 0;
        border-left: 4px solid #FFAB76;
        border-right: none;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        animation: slideInRight 0.4s ease;
        direction: ltr;
        text-align: left;
    }
    
    @keyframes slideInRight {
        from { opacity: 0; transform: translateX(20px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    @keyframes slideInLeft {
        from { opacity: 0; transform: translateX(-20px); }
        to { opacity: 1; transform: translateX(0); }
    }
    
    /* File upload - main area only */
    .main [data-testid="stFileUploader"] {
        background: white;
        border-radius: 16px;
        padding: 20px;
        border: 2px dashed #FFAB76;
    }
    
    .main [data-testid="stFileUploader"]:hover {
        border-color: #FF8C5A;
        background: #FFFBF5;
    }
    
    /* Text area / input */
    .stTextArea textarea, .stTextInput input {
        border-radius: 12px;
        border: 2px solid #FFE4D6;
        background: white;
        direction: ltr;
        text-align: left;
    }
    
    .stTextArea textarea:focus, .stTextInput input:focus {
        border-color: #FF8C5A;
        box-shadow: 0 0 15px rgba(255, 140, 90, 0.2);
    }
    
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
        background: white;
        border-radius: 12px;
        padding: 8px;
        direction: ltr !important;
        display: flex !important;
        flex-direction: row !important;
        justify-content: flex-start !important;
    }
    
    .stTabs [data-baseweb="tab"] {
        border-radius: 10px;
        padding: 12px 24px;
        background: #FFF5EB;
    }
    
    .stTabs [aria-selected="true"] {
        background: linear-gradient(90deg, #FF8C5A 0%, #FFAB76 100%) !important;
        color: white !important;
    }
    
    /* Progress bar */
    .stProgress > div > div {
        background: linear-gradient(90deg, #FF8C5A 0%, #FFAB76 100%);
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        background: linear-gradient(90deg, #FFF8F0 0%, #FFE8D6 100%);
        border-radius: 12px;
        font-weight: bold;
        direction: ltr;
        text-align: left;
    }
    
    .main [data-testid="stExpander"] {
        direction: ltr;
        text-align: left;
    }
    
    .main [data-testid="stExpander"] p,
    .main [data-testid="stExpander"] div {
        direction: ltr !important;
        text-align: left !important;
    }
    
    /* Logo */
    .logo-container {
        text-align: center;
        padding: 25px;
        background: linear-gradient(135deg, #FFF8F0 0%, #FFE8D6 100%);
        border-radius: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 15px rgba(255, 140, 90, 0.15);
    }
    
    .logo-title {
        font-size: 28px;
        font-weight: bold;
        color: #E07B39;
        margin: 0;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
        line-height: 1.4;
    }
    
    .logo-subtitle {
        font-size: 16px;
        color: #D4703A;
        margin-top: 10px;
    }
    
    /* Footer */
    .footer {
        text-align: center;
        padding: 25px;
        color: #D4703A;
        font-size: 14px;
        margin-top: 40px;
        border-top: 2px solid #FFE4D6;
        background: linear-gradient(180deg, transparent 0%, #FFF8F0 100%);
    }
    
    .designer-name {
        font-size: 16px;
        font-weight: bold;
        color: #E07B39;
        margin-top: 8px;
    }
</style>
""", unsafe_allow_html=True)

# Enhanced analysis instructions (system prompt)
SYSTEM_PROMPT = """
You are an expert specializing in explaining research papers in Econometrics. Your task is to analyze academic articles and provide a comprehensive, accurate, and in-depth explanation that is simplified and clear for researchers and graduate students.

## 🎯 Primary Goal:
Enable the researcher to fully understand the model/test: what it is, its conditions, assumptions, when to use it, when NOT to use it, special cases, and all related details.

## 📋 Comprehensive Analysis Methodology:

### 1️⃣ Structural Dissection of the Paper:

#### Introduction:
- Identify the **Main Contribution** precisely
- What problem does this paper solve?
- How does it differ from previous work?

#### Setup/Framework:
- Explain the basic **functional form** (e.g., Y = m(X) + U)
- Clarify the **properties of the error term** U
- Explain every **symbol** used in detail

#### Assumptions:
- Explain **each assumption** (A1, A2, A3...) in detail
- Clarify **why** each assumption was imposed
- Is it **realistic** in practice?
- What happens if the assumption is **violated**?

#### Theorems:
- Explain each theorem in simple language
- **Theorem 1**: usually Consistency
- **Theorem 2**: Asymptotic Normality
- What is the **practical implication** of each theorem?

### 2️⃣ Extracting Information from Specific Sections:

#### 🧪 From the Simulation / Monte Carlo Section:
- **Sample sizes** tested (N = 50, 100, 500, 1000...)
- **Number of replications**
- **Bias** in small samples
- **Standard deviation (SD)** and Root Mean Squared Error (RMSE)
- **Performance comparison** with alternative methods
- **Actual Size** vs. nominal size
- **Power** in different scenarios

#### 📝 From Footnotes:
- **Important technical notes**
- **Exceptions** and caveats
- **References** to key papers
- **Clarifications** of terminology

#### 💻 From the Empirical Application Section:
- **Data used** (source, period, variables)
- **How to apply** step by step
- **Interpreting results** and coefficients
- **Comparison** with other methods

#### 📎 From the Appendix:
- **Mathematical proofs** (summary)
- **Special cases** and generalizations
- **Additional tables** and sensitivity results

### 3️⃣ Required Response Structure:

---

## 📋 Executive Summary
[One comprehensive paragraph summarizing the paper and the model]

## 🎯 What is the Model/Test?
[Simple, clear explanation of the core idea]
- Full name and abbreviation
- Main purpose
- Historical context

## 📐 Mathematical Formulation
[Core equations]
- Main equation
- Explanation of each symbol
- Auxiliary equations (if any)

## ✅ Assumptions and Conditions
[Detailed list of all assumptions]
| Assumption | Description | Reason | What if violated |
|------------|-------------|--------|------------------|

## ⏰ When to Use This Model/Test?
[Appropriate usage cases in detail]
- Optimal conditions
- Suitable data types
- Problems it solves

## ⛔ When NOT to Use It?
[Inappropriate cases and limitations]
- Usage prohibitions
- Suggested alternatives for each case

## 🔄 Special Cases and Extensions
[Exceptions and variations]
- Small sample corrections
- Outlier data cases
- Possible generalizations

## 📈 Historical Development
[How this model evolved]
- Original paper and author(s)
- Subsequent developments
- Improved versions

## 🧪 Monte Carlo Simulation Summary
[Summary table from simulation section]
| Sample Size | Bias | RMSE | Actual Size | Power |
|-------------|------|------|-------------|-------|

**Interpretation**: [Explain the results]

## 💻 Practical Application Steps
[How to apply the model step by step]
1. Data preparation
2. Testing assumptions
3. Estimation
4. Inference

## 📚 Suggested Learning Plan
[What the researcher needs to learn before applying]

**Mathematical Requirements:**
- [ ] Matrix algebra
- [ ] Asymptotic theory
- [ ] ...

**Core Concepts:**
- [ ] ...

**Reference Papers to Read:**
1. [Paper] - [Reason]
2. ...

## 🔍 Identification Strategy
[How does the paper establish causal relationships?]
- Method used
- Instrumental variables (if any)
- Strengths and weaknesses

## 🛡️ Robustness Checks
[Summary of all robustness checks in the paper]
| Test | Result | Interpretation |
|------|--------|----------------|

## 📊 Main Table Analysis
[Detailed analysis of the key table in the paper]

## ⚠️ Econometric Issues
[How the paper handles potential problems]

## ✅ Strengths
[What the paper did well]

## ❌ Weaknesses and Limitations
[Limitations and caveats]

## 🔗 References and Resources
[Recommended papers and textbooks]

---

## 💡 Explanation Quality Rules:

1. **Accuracy**: Provide accurate and in-depth information directly from the paper
2. **Simplicity**: Explain complex concepts in plain language
3. **Comprehensiveness**: Don't leave out any important detail
4. **Organization**: Always use the specified structure
5. **Examples**: Provide illustrative examples when needed
6. **Toy Model Trick**: Start with the simple case (K=1), then generalize
7. **Concept Linking**: Link the model to similar or alternative models
8. **Warnings**: Alert to common mistakes in application
9. **Neutrality**: Provide objective criticism (strengths and weaknesses)
10. **Practicality**: Focus on what the researcher actually needs to apply

## ⚠️ Important Notes:
- If information is not in the paper, say so explicitly
- If something is ambiguous, point it out
- Offer alternatives when the model is inappropriate
- Use English technical terms alongside plain explanations

---

## 🔬 Advanced Instructions for Deeper Analysis:

### 4️⃣ Identification Strategy Analysis:
- What **identification strategy** is used? (IV, RDD, DID, Matching...)
- Are there **instrumental variables**? What are they?
- What are the **validity conditions** for the instruments? (Relevance, Exogeneity)
- Is the identification strategy **convincing**? What are the potential weaknesses?

### 5️⃣ Robustness Checks Analysis:
Find and extract all robustness checks in the paper:
- **Changing control variables**: Are results sensitive to adding/removing variables?
- **Different sub-samples**: Are results consistent across time periods/groups?
- **Alternative estimators**: (OLS, IV, GMM, Fixed Effects...)
- **Alternative variable definitions**: Do results change with different definitions?
- **Placebo tests**: Are there placebo tests? What did they show?
- **Sensitivity Analysis**: How sensitive are results to assumptions?

### 6️⃣ Deep Dive into Tables:
For each table in the paper:
- **Main table**: Identify core columns vs robustness columns
- **Stars and statistical significance**: (* p<0.10, ** p<0.05, *** p<0.01)
- **Standard errors**: Are they Clustered? Robust? At what level?
- **R² and Adjusted R²**: What is the model's explanatory power?
- **Number of observations (N)**: Does it change across columns? Why?
- **Fixed Effects used**: (Year FE, Industry FE, Firm FE...)

### 7️⃣ Econometric Issues:
Identify how the paper addresses:
- **Endogeneity**: Reverse causality or omitted variables?
- **Heteroskedasticity**: What solution is used?
- **Autocorrelation**: Especially in time-series data
- **Small Sample Bias**: Is N sufficient?
- **Selection Bias**: Is there sample selection bias?
- **Outliers**: How were they handled?

### 8️⃣ Reading Between the Lines:
- **Authors' tone**: Do they overstate or understate results?
- **What is unsaid**: Are there expected results not mentioned?
- **Potential researcher bias**: Do they have a prior position?
- **Gaps**: What does the paper leave unanswered?

### 9️⃣ Comprehensive Critical Evaluation:

#### Strengths:
- What did the paper do well?
- How convincing is the identification strategy?
- Are the data unique or valuable?

#### Weaknesses:
- What assumptions are unrealistic?
- What are threats to Internal Validity?
- What are threats to External Validity?
- Can the results be generalized?

#### Unanswered Questions:
- What questions did the paper leave open?
- What are possible research extensions?

### 🔟 Unified Table Reading Structure:

```
📊 Table [Number] Analysis:
├── Purpose: [What does this table test?]
├── Dependent Variable: [Y]
├── Main Independent Variables: [X1, X2...]
├── Control Variables: [Controls]
├── Fixed Effects: [Type]
├── Clustering: [At what level]
├── Main Result: [β = ?, significant?]
├── Economic Interpretation: [What does this mean in practice?]
└── Notes: [Anything unusual]
```

### 📖 Optimal Reading Strategy (5 Stages):

**Stage 1 – Quick Scan (5 minutes):**
- Title → Abstract → Introduction (first & last paragraph) → Conclusion
- Goal: Understand the question and main answer

**Stage 2 – Structure Understanding (10 minutes):**
- Browse tables and figures
- Read section headings
- Goal: Understand the argument structure

**Stage 3 – Selective Reading (20 minutes):**
- The Setup/Model section
- Data section
- Main Results table
- Goal: Understand the methodology and core results

**Stage 4 – Critical Reading (15 minutes):**
- Robustness section
- Important footnotes
- Goal: Evaluate the reliability of results

**Stage 5 – Deep Dive (as needed):**
- Appendix proofs
- Technical details
- Goal: Understand fine-grained details

### 🎓 Always Add at the End of the Analysis:

## 📝 Quick-Reference Summary
[One-page table summarizing everything]

| Element | Details |
|---------|---------|
| Model/Test Name | |
| Main Purpose | |
| Core Equation | |
| Critical Assumptions | |
| When to Use | |
| When NOT to Use | |
| Available Alternatives | |
| Software/Code | (Stata, R, Python) |

## ❓ Self-Review Questions
[5-10 questions the researcher can use to test their understanding]

## 🚨 Common Mistakes in Application
[List of mistakes to avoid]

## 📚 Suggested Reading Path
[Order of papers to read for deeper understanding]
1. Preliminary paper: [Paper name]
2. Current paper
3. Deep-dive paper: [Paper name]
4. Applied paper: [Paper name]

---

## 🔬 Additional Advanced Instructions:

### 1️⃣1️⃣ Estimator Properties Analysis:
When explaining any new estimator, clarify the following properties in detail:

#### a) Consistency:
- Is the estimator consistent? i.e., does plim(β̂) = β?
- What conditions are required for consistency?
- What happens if these conditions are violated?
- **Simple interpretation**: "As sample size grows, the estimator converges to the true value"

#### b) Unbiasedness:
- Does E(β̂) = β?
- If biased, how large is the bias?
- Does the bias vanish as sample size grows (Asymptotically Unbiased)?
- **Important note**: An estimator can be consistent yet biased in small samples

#### c) Efficiency:
- Is the estimator efficient? i.e., does it have the minimum possible variance?
- Does it achieve the Cramér-Rao lower bound?
- Compare variance with alternative estimators
- **BLUE**: Best Linear Unbiased Estimator (for OLS under Gauss-Markov conditions)

#### d) Asymptotic Normality:
- Does √n(β̂ − β) →d N(0, V)?
- What is the variance-covariance matrix V?
- How do we estimate V? (Robust, Clustered, HAC...)
- **Importance**: Allows constructing confidence intervals and hypothesis tests

### 1️⃣2️⃣ Monte Carlo Deep Dive:

#### Extracting Information from Simulation Tables:
```
📊 Simulation Analysis Table:
┌─────────────────────────────────────────────────────────────┐
│ Metric         │ Definition                  │ Ideal Value  │
├─────────────────────────────────────────────────────────────┤
│ Bias           │ E(β̂) − β                   │ ≈ 0          │
│ Std. Dev (SD)  │ Standard deviation of est.  │ Small        │
│ RMSE           │ √(Bias² + Variance)         │ Small        │
│ Size (5%)      │ Rejection rate under H₀     │ ≈ 0.05       │
│ Power          │ Rejection rate under H₁     │ Close to 1   │
│ Coverage       │ CI coverage of true value   │ ≈ 0.95       │
└─────────────────────────────────────────────────────────────┘
```

#### Questions to Answer from the Simulation:
1. **How many replications?** (1000, 5000, 10000...) — more is better
2. **What sample sizes are tested?** (N = 50, 100, 200, 500, 1000...)
3. **How does the estimator behave in small vs. large samples?**
4. **Is Size close to the nominal level (5%)?** — if not, the test is distorted
5. **How does Power change with sample size and effect size?**
6. **What DGP (Data Generating Process) is used?** — is it realistic?

#### Interpreting Results:
| Observation | Interpretation |
|-------------|----------------|
| Large Bias with small N, vanishes with large N | Small-sample bias; estimator is consistent |
| Size > 0.05 | Test over-rejects |
| Size < 0.05 | Test is too conservative |
| Low Power | Test is weak at detecting the alternative |
| RMSE does not shrink with N | Consistency problem |

### 1️⃣3️⃣ Interpreting Regression Coefficients and Standard Errors:

#### Reading a Regression Table:
```
┌────────────────────────────────────────────────────────┐
│ Variable │ Coef (β̂) │ Std.Err │ t-stat │ P>|t| │ CI  │
├────────────────────────────────────────────────────────┤
│ X₁       │ 0.523    │ 0.102   │ 5.13   │ 0.000 │ ... │
│ X₂       │ -0.087   │ 0.045   │ -1.93  │ 0.054 │ ... │
└────────────────────────────────────────────────────────┘

Interpretation:
• β̂ = 0.523: A one-unit increase in X₁ is associated with a 0.523 increase in Y
• SE = 0.102: Precision of the estimate (smaller is better)
• t = β̂/SE = 5.13: Is the coefficient significantly different from zero?
• p-value = 0.000: Probability of obtaining |t| ≥ 5.13 under H₀
• CI: 95% confidence interval for the true coefficient
```

#### Types of Standard Errors:
| Type | When to Use | Stata Code |
|------|-------------|------------|
| Classical | Homoskedastic, independent | reg y x |
| Robust (HC) | Heteroskedasticity | reg y x, robust |
| Clustered | Correlation within groups | reg y x, cluster(id) |
| HAC (Newey-West) | Autocorrelation in time series | newey y x, lag(4) |
| Two-way Clustered | Correlation along two dimensions | reghdfe y x, cluster(firm year) |

### 1️⃣4️⃣ Statistical vs. Economic Significance:

#### Statistical Significance:
- p < 0.05 → "statistically significant"
- But this does NOT necessarily mean the effect is "important"!

#### Economic Significance:
- How large is the effect? Is it practically meaningful?
- Use the Standardized Effect
- Compare with results from prior literature
- **Example**: β = 0.001*** (highly significant but economically negligible)

#### Computing Economic Significance:
```
Standardized Effect = β × (SD_X / SD_Y)
Interpretation: A one-standard-deviation increase in X is associated
with a [value] standard-deviation change in Y
```

### 1️⃣5️⃣ IV Diagnostics:

#### If the paper uses Instrumental Variables:
1. **Relevance test**: F-statistic in the First Stage
   - F > 10 → Strong instrument
   - F < 10 → Weak instruments problem

2. **Exogeneity test**: Cannot be directly tested!
   - Relies on theoretical argument
   - Overidentification test (if # instruments > # endogenous variables)

3. **Tests to look for**:
   - Cragg-Donald F-statistic
   - Kleibergen-Paap rk Wald F-statistic (for Robust SE)
   - Hansen J-statistic (Overidentification)
   - Anderson-Rubin test (Weak-IV robust)

### 1️⃣6️⃣ Panel Data Specifics:

#### If the paper uses Panel Data:
| Estimator | Core Assumption | When to Use |
|-----------|-----------------|-------------|
| Pooled OLS | No individual effect | Rarely appropriate |
| Fixed Effects (FE) | Corr(αᵢ, Xᵢₜ) ≠ 0 | Individual effect correlated with regressors |
| Random Effects (RE) | Corr(αᵢ, Xᵢₜ) = 0 | Individual effect independent |
| First Differences | Same as FE | Alternative to FE |

#### Hausman Test:
- H₀: RE is consistent and efficient
- H₁: FE is consistent, RE is not
- If H₀ rejected → Use FE

### 1️⃣7️⃣ Reading Figures:

#### Common Figure Types:
1. **Scatter Plot + Regression Line**: Bivariate relationship
2. **Residual Plots**: Testing assumptions
3. **Event Study Graph**: Dynamic effects around an event
4. **RDD Plot**: Discontinuity at the threshold
5. **Binned Scatter**: Non-linear relationship

#### What to Look For:
- Is the relationship linear or curved?
- Are there influential outliers?
- Are the confidence intervals tight or wide?
- Is the pattern consistent with quantitative results?

### 1️⃣8️⃣ Code and Software Application:

#### If the paper provides code or mentions software:
```
📦 Application Info:
├── Software: [Stata / R / Python / Matlab]
├── Packages used: [reghdfe, fixest, linearmodels...]
├── Main commands: [...]
├── Replication Files: [link if available]
└── Data: [Available / Not available]
```

#### Common Stata Commands in Papers:
| Command | Usage |
|---------|-------|
| reg | Simple OLS |
| xtreg, fe | Fixed Effects |
| ivregress 2sls | Instrumental Variables |
| reghdfe | High-dimensional Fixed Effects |
| didregress | Difference-in-Differences |
| rdrobust | Regression Discontinuity |

### 1️⃣9️⃣ Final Checklist:

Before finishing the analysis, make sure to answer:

- [ ] What is the main research question?
- [ ] What is the contribution relative to the literature?
- [ ] What model/test is proposed?
- [ ] What assumptions are required?
- [ ] What is the identification strategy?
- [ ] What data are used?
- [ ] What are the main results?
- [ ] Are the results robust?
- [ ] What are the limitations and caveats?
- [ ] What are the practical/policy implications?
- [ ] When should this model be used?
- [ ] When should it NOT be used?
- [ ] What alternatives are available?
- [ ] What reference papers are recommended for reading?

### 2️⃣0️⃣ Tips for the Beginner Researcher:

1. **Don't fear the math** — focus on intuition first
2. **Start with the simple case** — K=1, no Fixed Effects
3. **Read the explanatory paper first** — if one exists
4. **Don't read the proofs in the first pass**
5. **Build your own symbol dictionary**
6. **Compare with what you know** — how does this differ from standard OLS?
7. **Ask: Why this method and not another?**
8. **Learn from mistakes** — the first paper is hard, the fifth is easier
"""

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'analysis_result' not in st.session_state:
    st.session_state.analysis_result = None
if 'uploaded_files_content' not in st.session_state:
    st.session_state.uploaded_files_content = {}
if 'files_names' not in st.session_state:
    st.session_state.files_names = []
if 'client' not in st.session_state:
    st.session_state.client = None
if 'thinking_level' not in st.session_state:
    st.session_state.thinking_level = "HIGH"
if 'saved_api_key' not in st.session_state:
    st.session_state.saved_api_key = ""
if 'api_key_saved' not in st.session_state:
    st.session_state.api_key_saved = False


def create_word_report(content: str, title: str = "Analysis Report") -> bytes:
    """Create a Word file from content with LTR English support"""
    doc = DocxDocument()
    
    # Page setup (LTR)
    section = doc.sections[0]
    
    # Main title
    title_para = doc.add_heading("📊 Econometric Paper Analysis Platform", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Report info
    date_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    info_para = doc.add_paragraph()
    info_para.add_run(f"Date: {date_str}").bold = True
    info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    paper_para = doc.add_paragraph()
    paper_para.add_run(f"Article: {title}").bold = True
    paper_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # empty line
    
    # Process content
    lines = content.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('## '):
                heading_text = line.replace('## ', '').replace('#', '').strip()
                heading = doc.add_heading(heading_text, level=2)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('### '):
                heading_text = line.replace('### ', '').replace('#', '').strip()
                heading = doc.add_heading(heading_text, level=3)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('# '):
                heading_text = line.replace('# ', '').strip()
                heading = doc.add_heading(heading_text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif line.startswith('|') or line.startswith('---'):
                # Skip Markdown tables
                continue
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet lists
                bullet_text = line.replace('- ', '').replace('* ', '').strip()
                para = doc.add_paragraph(bullet_text, style='List Bullet')
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # Normal text
                clean_line = line.replace('**', '').replace('*', '').replace('`', '')
                clean_line = clean_line.replace('###', '').replace('##', '').replace('#', '')
                if clean_line.strip():
                    para = doc.add_paragraph(clean_line.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph()
    footer.add_run("Design & Development: Dr. Merwan Roudane")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document to memory
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def upload_file_to_gemini(client, file_content: bytes, filename: str):
    """Upload a file to Gemini using the new SDK"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
        tmp.write(file_content)
        tmp_path = tmp.name
    
    try:
        uploaded_file = client.files.upload(file=tmp_path, config={"display_name": filename})
        return uploaded_file
    finally:
        os.unlink(tmp_path)


def analyze_with_gemini(client, files_content: dict, thinking_level: str, question: str = None) -> str:
    """Analyze articles using Gemini 3 Flash"""
    
    # Upload files
    uploaded_files = []
    for filename, file_data in files_content.items():
        try:
            uploaded_file = upload_file_to_gemini(client, file_data, filename)
            uploaded_files.append(uploaded_file)
        except Exception as e:
            st.warning(f"Could not upload {filename}: {str(e)}")
    
    if not uploaded_files:
        return "❌ No files were uploaded successfully"
    
    # Build prompt
    if question:
        prompt = f"""
{SYSTEM_PROMPT}

Based on the attached articles, answer the following question in detail and accurately:

{question}

Provide a comprehensive answer referencing the relevant sections of the article.
"""
    else:
        prompt = f"""
{SYSTEM_PROMPT}

Perform a comprehensive and in-depth analysis of the attached articles following the specified structure and instructions.

Focus on:
1. Explaining the main model/test in a simple and clear way
2. Extracting all assumptions and explaining them
3. Identifying precisely when to use and when NOT to use the model
4. Extracting simulation information and organizing it in a table
5. Providing a practical learning plan for the researcher
6. Pointing out special cases and warnings

Goal: Enable the researcher to fully understand the model and apply it in practice.
"""
    
    try:
        # Set thinking level
        thinking_level_map = {
            "HIGH": types.ThinkingLevel.HIGH,
            "MEDIUM": types.ThinkingLevel.MEDIUM,
            "LOW": types.ThinkingLevel.LOW,
            "MINIMAL": types.ThinkingLevel.MINIMAL,
        }
        level = thinking_level_map.get(thinking_level, types.ThinkingLevel.HIGH)
        
        # Build content parts
        content_parts = uploaded_files + [prompt]
        
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=content_parts,
            config=types.GenerateContentConfig(
                thinking_config=types.ThinkingConfig(
                    thinking_level=level
                )
            )
        )
        return response.text
    except Exception as e:
        return f"❌ An error occurred during analysis: {str(e)}"


def chat_with_gemini(client, files_content: dict, thinking_level: str, messages: list, new_message: str) -> str:
    """Chat about articles using Gemini 3 Flash"""
    
    # Upload files
    uploaded_files = []
    for filename, file_data in files_content.items():
        try:
            uploaded_file = upload_file_to_gemini(client, file_data, filename)
            uploaded_files.append(uploaded_file)
        except:
            pass
    
    # Build conversation context
    chat_context = "\n\n".join([
        f"{'User' if m['role'] == 'user' else 'Assistant'}: {m['content']}"
        for m in messages[-6:]
    ])
    
    prompt = f"""
{SYSTEM_PROMPT}

Based on the attached articles and the previous conversation:

{chat_context}

New question: {new_message}

Answer accurately and in detail, referencing the relevant parts of the article. If the question is about something unclear, explain it in a simple way with examples.
"""
    
    try:
        # Set thinking level
        thinking_level_map = {
            "HIGH": types.ThinkingLevel.HIGH,
            "MEDIUM": types.ThinkingLevel.MEDIUM,
            "LOW": types.ThinkingLevel.LOW,
            "MINIMAL": types.ThinkingLevel.MINIMAL,
        }
        level = thinking_level_map.get(thinking_level, types.ThinkingLevel.HIGH)
        
        content_parts = uploaded_files + [prompt]
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=content_parts,
            config=types.GenerateContentConfig(
                thinking_config=types.ThinkingConfig(
                    thinking_level=level
                )
            )
        )
        return response.text
    except Exception as e:
        return f"❌ An error occurred: {str(e)}"


# ==================== Main Interface ====================

# Sidebar
with st.sidebar:
    st.markdown("""
    <div style='text-align: center; padding: 15px;'>
        <h3 style='color: white; margin: 0; font-size: 16px;'>📊 Comprehensive Platform</h3>
        <p style='color: #FFE4D6; font-size: 11px;'>Econometrics</p>
    </div>
    """, unsafe_allow_html=True)
    
    st.markdown("---")
    
    st.markdown("### 🔑 API Key")
    
    # Check for saved key
    if st.session_state.api_key_saved and st.session_state.saved_api_key:
        st.success("✅ Key saved")
        st.markdown(f"<p style='color: #FFE4D6; font-size: 12px;'>Key: ****{st.session_state.saved_api_key[-4:]}</p>", unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🔄 Change", use_container_width=True, key="change_key"):
                st.session_state.api_key_saved = False
                st.session_state.saved_api_key = ""
                st.session_state.client = None
                st.rerun()
        
        api_key = st.session_state.saved_api_key
    else:
        api_key = st.text_input(
            "Enter your Google API Key",
            type="password",
            help="Get your API key from Google AI Studio",
            label_visibility="collapsed",
            placeholder="Enter API key here..."
        )
        
        if api_key:
            col1, col2 = st.columns(2)
            with col1:
                if st.button("💾 Save", use_container_width=True, key="save_key"):
                    st.session_state.saved_api_key = api_key
                    st.session_state.api_key_saved = True
                    st.success("✅ Saved!")
                    st.rerun()
    
    st.markdown("### ⚙️ Model Settings")
    thinking_level = st.selectbox(
        "Thinking Level",
        ["HIGH", "MEDIUM", "LOW", "MINIMAL"],
        index=0,
        help="HIGH: Deep thinking (default) | MEDIUM: Balanced | LOW: Fast | MINIMAL: Fastest"
    )
    
    if api_key:
        if not st.session_state.api_key_saved:
            st.success("✅ Key entered")
        try:
            # Create client using new SDK
            client = genai.Client(api_key=api_key)
            st.session_state.client = client
            st.session_state.thinking_level = thinking_level
        except Exception as e:
            st.error(f"Model configuration error: {str(e)}")
    else:
        st.warning("⚠️ API key required")
    
    st.markdown("---")
    
    st.markdown("### 📁 Upload Articles")
    uploaded_files = st.file_uploader(
        "Select PDF files",
        type=['pdf'],
        accept_multiple_files=True,
        help="Maximum: 10 articles",
        label_visibility="collapsed"
    )
    
    if uploaded_files:
        if len(uploaded_files) > 10:
            st.error("⚠️ Maximum is 10 files!")
            uploaded_files = uploaded_files[:10]
        
        st.success(f"✅ {len(uploaded_files)} file(s) uploaded")
        
        st.session_state.uploaded_files_content = {}
        st.session_state.files_names = []
        for file in uploaded_files:
            content = file.read()
            st.session_state.uploaded_files_content[file.name] = content
            st.session_state.files_names.append(file.name)
            file.seek(0)
        
        with st.expander("📋 Uploaded Files", expanded=False):
            for name in st.session_state.files_names:
                st.write(f"📄 {name}")
    
    st.markdown("---")
    
    st.markdown("""
    <div style='text-align: center; padding: 15px; background: rgba(255,255,255,0.1); border-radius: 10px;'>
        <p style='font-size: 11px; margin: 0; color: #FFE4D6;'>Design & Development</p>
        <p style='font-size: 14px; font-weight: bold; margin: 5px 0; color: white;'>Dr. Merwan Roudane</p>

    </div>
    """, unsafe_allow_html=True)


# Main content area
st.markdown("""
<div class='logo-container'>
    <p class='logo-title'>📊 Comprehensive Platform for Explaining Original Econometric Papers and Tests</p>
    <p class='logo-subtitle'>Design & Development: Dr. Merwan Roudane</p>
</div>
""", unsafe_allow_html=True)

# Tabs
tab1, tab2, tab3 = st.tabs(["📰 Article Analysis", "💬 Chat Window", "📥 Download Report"])

# ==================== Analysis Tab ====================
with tab1:
    st.markdown("### 🔍 Comprehensive Article Analysis")
    
    if not api_key:
        st.warning("⚠️ Please enter your Google API key in the sidebar to get started")
    elif not st.session_state.uploaded_files_content:
        st.info("📁 Please upload at least one article from the sidebar")
    else:
        col1, col2 = st.columns([3, 1])
        
        with col1:
            st.markdown(f"""
            <div class='info-card'>
                <h4 style='color: #E07B39; margin: 0;'>📚 Articles Ready for Analysis</h4>
                <p style='font-size: 24px; font-weight: bold; color: #D4703A; margin: 10px 0;'>{len(st.session_state.uploaded_files_content)}</p>
                <p style='color: #888; font-size: 13px;'>Click "Start Analysis" to get a comprehensive and detailed explanation</p>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            analyze_btn = st.button("🚀 Start Analysis", use_container_width=True, type="primary")
        
        if analyze_btn:
            if st.session_state.client is None:
                st.error("❌ Please make sure you have entered a valid API key")
            else:
                with st.spinner("⏳ Analyzing articles... This may take a few minutes"):
                    try:
                        result = analyze_with_gemini(
                            st.session_state.client,
                            st.session_state.uploaded_files_content,
                            st.session_state.thinking_level
                        )
                        st.session_state.analysis_result = result
                    except Exception as e:
                        st.error(f"❌ An error occurred: {str(e)}")
        
        if st.session_state.analysis_result:
            st.markdown("---")
            st.markdown("### 📊 Analysis Results")
            
            with st.container():
                st.markdown(f"""
                <div style='background: white; padding: 25px; border-radius: 15px; 
                            box-shadow: 0 4px 15px rgba(0,0,0,0.08); border: 1px solid #FFE4D6;'>
                """, unsafe_allow_html=True)
                
                st.markdown(st.session_state.analysis_result)
                
                st.markdown("</div>", unsafe_allow_html=True)

# ==================== Chat Tab ====================
with tab2:
    st.markdown("### 💬 Interactive Chat Window")
    st.markdown("Ask your questions about the uploaded articles to get additional clarifications")
    
    if not api_key:
        st.warning("⚠️ Please enter your API key first")
    elif not st.session_state.uploaded_files_content:
        st.info("📁 Please upload at least one article")
    else:
        chat_container = st.container()
        
        with chat_container:
            for msg in st.session_state.messages:
                if msg['role'] == 'user':
                    st.markdown(f"""
                    <div class='chat-user'>
                        <strong>🧑‍🎓 You:</strong><br>{msg['content']}
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div class='chat-assistant'>
                        <strong>🤖 Platform:</strong><br>{msg['content']}
                    </div>
                    """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        col1, col2 = st.columns([5, 1])
        
        with col1:
            user_question = st.text_area(
                "Type your question here",
                placeholder="Example: What is Assumption A3 in detail? Or: Explain the simulation results...",
                height=80,
                label_visibility="collapsed"
            )
        
        with col2:
            send_btn = st.button("📤 Send", use_container_width=True, type="primary")
        
        st.markdown("**💡 Suggested Questions:**")
        suggested_cols = st.columns(3)
        
        suggestions = [
            "What are the core assumptions?",
            "Explain the simulation results",
            "When should this model NOT be used?"
        ]
        
        for i, suggestion in enumerate(suggestions):
            with suggested_cols[i]:
                if st.button(suggestion, key=f"sug_{i}", use_container_width=True):
                    user_question = suggestion
                    send_btn = True
        
        if send_btn and user_question:
            if st.session_state.client is None:
                st.error("❌ Please make sure you have entered a valid API key")
            else:
                st.session_state.messages.append({
                    'role': 'user',
                    'content': user_question
                })
                
                with st.spinner("⏳ Thinking..."):
                    try:
                        response = chat_with_gemini(
                            st.session_state.client,
                            st.session_state.uploaded_files_content,
                            st.session_state.thinking_level,
                            st.session_state.messages,
                            user_question
                        )
                        
                        st.session_state.messages.append({
                            'role': 'assistant',
                            'content': response
                        })
                        
                        st.rerun()
                    except Exception as e:
                        st.error(f"❌ An error occurred: {str(e)}")
        
        if st.session_state.messages:
            if st.button("🗑️ Clear Conversation", type="secondary"):
                st.session_state.messages = []
                st.rerun()

# ==================== Download Tab ====================
with tab3:
    st.markdown("### 📥 Download Report")
    
    if st.session_state.analysis_result:
        st.success("✅ Report is ready for download!")
        
        st.markdown(f"""
        <div class='info-card'>
            <h4 style='color: #E07B39;'>📄 Report Information</h4>
            <p><strong>Number of articles analyzed:</strong> {len(st.session_state.files_names)}</p>
            <p><strong>Analysis date:</strong> {datetime.now().strftime("%Y-%m-%d %H:%M")}</p>
            <p><strong>Articles:</strong></p>
            <ul>
                {"".join([f"<li>{name}</li>" for name in st.session_state.files_names])}
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2 = st.columns(2)
        
        with col1:
            try:
                word_bytes = create_word_report(
                    st.session_state.analysis_result,
                    ", ".join(st.session_state.files_names[:2]) + ("..." if len(st.session_state.files_names) > 2 else "")
                )
                
                st.download_button(
                    label="📥 Download Word",
                    data=word_bytes,
                    file_name=f"econometrics_report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary"
                )
            except Exception as e:
                st.error(f"Error creating Word file: {str(e)}")
        
        with col2:
            st.download_button(
                label="📄 Download TXT",
                data=st.session_state.analysis_result,
                file_name=f"econometrics_report_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with st.expander("👁️ Preview Report", expanded=False):
            st.markdown(st.session_state.analysis_result)
    
    else:
        st.info("📊 Please analyze your articles first from the 'Article Analysis' tab to get a downloadable report")
        
        st.markdown("""
        <div class='info-card' style='text-align: center;'>
            <p style='font-size: 48px; margin: 20px 0;'>📑</p>
            <h4 style='color: #D4703A;'>No report yet</h4>
            <p style='color: #888;'>Upload your articles and analyze them to get a comprehensive report</p>
        </div>
        """, unsafe_allow_html=True)

# ==================== Footer ====================
st.markdown("---")
st.markdown("""
<div class='footer'>
    <p style='font-size: 15px; color: #E07B39; font-weight: bold;'>📊 Comprehensive Platform for Explaining Original Econometric Papers and Tests</p>
    <p class='designer-name'>Design & Development: Dr. Merwan Roudane</p>
    <p style='font-size: 11px; color: #999; margin-top: 15px;'>© 2025</p>
</div>
""", unsafe_allow_html=True)
